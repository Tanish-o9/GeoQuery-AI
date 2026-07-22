import logging
import io
import os
from typing import Dict, Any
from datetime import datetime

# ReportLab imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

# python-docx imports
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class ReportGeneratorService:
    """Service to automatically compile GIS spatial statistics and summaries into PDF and DOCX formats"""

    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Configure professional styles for the PDF report"""
        # Title
        self.title_style = ParagraphStyle(
            'ReportTitle',
            parent=self.styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=24,
            leading=28,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=20
        )
        
        # Subtitle
        self.subtitle_style = ParagraphStyle(
            'ReportSubtitle',
            parent=self.styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=12,
            leading=16,
            textColor=colors.HexColor('#4b5563'),
            spaceAfter=30
        )
        
        # Section Headings
        self.heading_style = ParagraphStyle(
            'ReportHeading',
            parent=self.styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor('#0284c7'),
            spaceBefore=15,
            spaceAfter=10,
            keepWithNext=True
        )
        
        # Normal Body Text
        self.body_style = ParagraphStyle(
            'ReportBody',
            parent=self.styles['BodyText'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#374151'),
            spaceAfter=10
        )
        
        # Risk Callout Box Text
        self.risk_style = ParagraphStyle(
            'ReportRisk',
            parent=self.styles['BodyText'],
            fontName='Helvetica-Bold',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#991b1b'),
            spaceAfter=10
        )

    def generate_pdf(self, data: Dict[str, Any]) -> io.BytesIO:
        """Generate PDF report inside memory stream"""
        logger.info("Generating spatial analysis PDF report...")
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54
        )
        
        story = []
        
        # 1. Title / Header Block
        story.append(Spacer(1, 10))
        story.append(Paragraph("GeoQuery AI - Geospatial Intelligence Report", self.title_style))
        date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        story.append(Paragraph(f"Generated on {date_str} UTC | Bounding Centroid: {data.get('centroid', {}).get('latitude')}, {data.get('centroid', {}).get('longitude')}", self.subtitle_style))
        story.append(Spacer(1, 10))
        
        # 2. Executive Summary
        story.append(Paragraph("1. Executive Summary", self.heading_style))
        summary_text = (
            f"This geospatial report evaluates the geographic, ecological, and infrastructural properties "
            f"of the designated Region of Interest (ROI) located around coordinates "
            f"({data.get('centroid', {}).get('latitude')}, {data.get('centroid', {}).get('longitude')}). "
            f"The analyzed boundary covers {data.get('area', {}).get('hectares')} hectares ({data.get('area', {}).get('sq_km')} sq km) "
            f"with a boundary perimeter measuring {data.get('perimeter_m')} meters. "
            f"The environment is categorized as having a flood risk level of '{data.get('flood_risk', {}).get('level')}' "
            f"and supports an estimated local resident population of {data.get('population_estimation') or 0} individuals."
        )
        story.append(Paragraph(summary_text, self.body_style))
        story.append(Spacer(1, 10))
        
        # 3. Spatial Statistics Table
        story.append(Paragraph("2. Primary Spatial Metrics", self.heading_style))
        table_data = [
            [Paragraph("<b>Metric</b>", self.body_style), Paragraph("<b>Value</b>", self.body_style), Paragraph("<b>Description</b>", self.body_style)],
            ["Total Area (Hectares)", f"{data.get('area', {}).get('hectares')}", "Land cover area in hectares"],
            ["Total Area (Sq Km)", f"{data.get('area', {}).get('sq_km')}", "Land cover area in square kilometers"],
            ["Perimeter (meters)", f"{data.get('perimeter_m')}", "Total boundary path length"],
            ["Est. Population", f"{data.get('population_estimation')}", "Calculated proxy settlement occupancy"],
            ["Flood Risk Class", f"{data.get('flood_risk', {}).get('level')}", "Historical and proximity flood proxy category"],
        ]
        
        t = Table(table_data, colWidths=[2.0*inch, 1.5*inch, 3.5*inch])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f3f4f6')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('BOTTOMPADDING', (0,0), (-1,0), 8),
            ('TOPPADDING', (0,0), (-1,0), 8),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')]),
            ('FONTNAME', (0,0), (-1,-1), 'Helvetica'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        story.append(t)
        story.append(Spacer(1, 15))
        
        # 4. Land Use Assessment
        story.append(Paragraph("3. Land Use & Surface Cover Composition", self.heading_style))
        lu = data.get('land_use', {})
        lu_data = [
            [Paragraph("<b>Land Use Class</b>", self.body_style), Paragraph("<b>Percentage (%)</b>", self.body_style), Paragraph("<b>Visual Representation Proxy</b>", self.body_style)],
            ["Urban / Built-up", f"{lu.get('urban', 0)}%", "Residential buildings, pavement, concrete"],
            ["Vegetation Canopy", f"{lu.get('vegetation', 0)}%", "Wooded zones, canopy, grass cover"],
            ["Agricultural Fields", f"{lu.get('agriculture', 0)}%", "Crops, agricultural fields, bare farms"],
            ["Surface Water", f"{lu.get('water', 0)}%", "Rivers, reservoirs, seasonal pools"],
            ["Bare Soil / Desert", f"{lu.get('bare_soil', 0)}%", "Exposed rocks, sand, uncultivated soils"]
        ]
        lu_table = Table(lu_data, colWidths=[2.5*inch, 1.5*inch, 3.0*inch])
        lu_table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f3f4f6')),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e5e7eb')),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f9fafb')]),
            ('BOTTOMPADDING', (0,0), (-1,-1), 6),
            ('TOPPADDING', (0,0), (-1,-1), 6),
        ]))
        story.append(lu_table)
        story.append(Spacer(1, 15))
        
        # 5. Nearby Amenities
        story.append(Paragraph("4. Infrastructural & Hydrological Amenities", self.heading_style))
        amenities = data.get('amenities', {})
        
        # Summarize roads, rivers, hospitals
        hosp_list = [f"{h.get('name')} (Distance: {h.get('distance_m')}m, Beds: {h.get('beds')})" for h in amenities.get('hospitals', [])[:3]]
        school_list = [f"{s.get('name')} (Distance: {s.get('distance_m')}m, Type: {s.get('type')})" for s in amenities.get('schools', [])[:3]]
        road_list = [f"{r.get('name')} (Distance: {r.get('distance_m')}m, Type: {r.get('type')})" for r in amenities.get('roads', [])[:3]]
        
        story.append(Paragraph(f"<b>Nearest Healthcare Centers:</b> {', '.join(hosp_list) if hosp_list else 'None detected'}.", self.body_style))
        story.append(Paragraph(f"<b>Nearest Educational Centers:</b> {', '.join(school_list) if school_list else 'None detected'}.", self.body_style))
        story.append(Paragraph(f"<b>Major Adjacent Roadways:</b> {', '.join(road_list) if road_list else 'None detected'}.", self.body_style))
        story.append(Spacer(1, 10))
        
        # 6. Risk Analysis & Recommendations
        story.append(Paragraph("5. Environmental Risk Analysis & Recommendations", self.heading_style))
        
        risk_level = data.get('flood_risk', {}).get('level', 'Low')
        if risk_level == "High":
            rec_text = (
                "CRITICAL WARNING: The region demonstrates high vulnerability to hydrological emergencies. "
                "It is strongly recommended to restrict permanent concrete civil structures in buffers closer than 500m "
                "to surface channels. Implement permeable pavement surfaces, sustainable drainage channels, "
                "and deploy real-time telemetry river monitors to build resilient local infrastructure."
            )
            story.append(Paragraph(rec_text, self.risk_style))
        else:
            rec_text = (
                "The analysis indicates low-to-moderate environmental hazards. Standard structural guidelines are "
                "adequate. Urban growth should prioritize preserving green pathways and planting vegetation strips "
                "to prevent heat island effects and retain historical natural absorption thresholds."
            )
            story.append(Paragraph(rec_text, self.body_style))
            
        doc.build(story)
        buffer.seek(0)
        return buffer

    def generate_docx(self, data: Dict[str, Any]) -> io.BytesIO:
        """Generate Microsoft Word DOCX report inside memory stream"""
        logger.info("Generating spatial analysis DOCX report...")
        doc = docx.Document()
        
        # Set Page margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("GeoQuery AI - Geospatial Intelligence Report")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(22)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Subtitle
        subtitle = doc.add_paragraph()
        date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        subtitle_run = subtitle.add_run(f"Generated: {date_str} UTC\nCentroid: {data.get('centroid', {}).get('latitude')}, {data.get('centroid', {}).get('longitude')}")
        subtitle_run.font.size = Pt(10.5)
        subtitle_run.font.italic = True
        
        # Section 1: Executive Summary
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("1. Executive Summary")
        h1_run.bold = True
        h1_run.font.size = Pt(14)
        h1_run.font.color.rgb = docx.shared.RGBColor(2, 132, 199)
        
        summary_text = (
            f"This geospatial report evaluates the geographic, ecological, and infrastructural properties "
            f"of the designated Region of Interest (ROI). The analyzed boundary covers "
            f"{data.get('area', {}).get('hectares')} hectares ({data.get('area', {}).get('sq_km')} sq km) "
            f"with a boundary perimeter measuring {data.get('perimeter_m')} meters. "
            f"The environment is categorized as having a flood risk level of '{data.get('flood_risk', {}).get('level')}' "
            f"and supports an estimated local resident population of {data.get('population_estimation') or 0} individuals."
        )
        doc.add_paragraph(summary_text)
        
        # Section 2: Primary Spatial Metrics Table
        h2 = doc.add_paragraph()
        h2_run = h2.add_run("2. Primary Spatial Metrics")
        h2_run.bold = True
        h2_run.font.size = Pt(14)
        h2_run.font.color.rgb = docx.shared.RGBColor(2, 132, 199)
        
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Light Shading Accent 1'
        
        headers = ["Metric", "Value", "Description"]
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = header
            hdr_cells[idx].paragraphs[0].runs[0].font.bold = True
            
        rows_data = [
            ("Total Area (Hectares)", f"{data.get('area', {}).get('hectares')}", "Land cover area in hectares"),
            ("Total Area (Sq Km)", f"{data.get('area', {}).get('sq_km')}", "Land cover area in square kilometers"),
            ("Perimeter (meters)", f"{data.get('perimeter_m')}", "Total boundary path length"),
            ("Est. Population", f"{data.get('population_estimation')}", "Calculated proxy settlement occupancy"),
            ("Flood Risk Class", f"{data.get('flood_risk', {}).get('level')}", "Historical and proximity flood proxy category"),
        ]
        
        for r_idx, row_data in enumerate(rows_data, start=1):
            row_cells = table.rows[r_idx].cells
            row_cells[0].text = row_data[0]
            row_cells[1].text = row_data[1]
            row_cells[2].text = row_data[2]
            
        # Section 3: Land Use Table
        doc.add_paragraph() # Spacer
        h3 = doc.add_paragraph()
        h3_run = h3.add_run("3. Land Use & Surface Cover Composition")
        h3_run.bold = True
        h3_run.font.size = Pt(14)
        h3_run.font.color.rgb = docx.shared.RGBColor(2, 132, 199)
        
        lu = data.get('land_use', {})
        lu_table = doc.add_table(rows=6, cols=2)
        lu_table.style = 'Light Shading Accent 1'
        
        lu_hdr = lu_table.rows[0].cells
        lu_hdr[0].text = "Land Use Class"
        lu_hdr[0].paragraphs[0].runs[0].font.bold = True
        lu_hdr[1].text = "Percentage (%)"
        lu_hdr[1].paragraphs[0].runs[0].font.bold = True
        
        lu_rows = [
            ("Urban / Built-up", f"{lu.get('urban', 0)}%"),
            ("Vegetation Canopy", f"{lu.get('vegetation', 0)}%"),
            ("Agricultural Fields", f"{lu.get('agriculture', 0)}%"),
            ("Surface Water", f"{lu.get('water', 0)}%"),
            ("Bare Soil / Desert", f"{lu.get('bare_soil', 0)}%")
        ]
        
        for r_idx, row_data in enumerate(lu_rows, start=1):
            row_cells = lu_table.rows[r_idx].cells
            row_cells[0].text = row_data[0]
            row_cells[1].text = row_data[1]
            
        # Section 4: Recommendations
        doc.add_paragraph() # Spacer
        h4 = doc.add_paragraph()
        h4_run = h4.add_run("4. Environmental Risk Analysis & Recommendations")
        h4_run.bold = True
        h4_run.font.size = Pt(14)
        h4_run.font.color.rgb = docx.shared.RGBColor(2, 132, 199)
        
        risk_level = data.get('flood_risk', {}).get('level', 'Low')
        p_rec = doc.add_paragraph()
        if risk_level == "High":
            p_rec_run = p_rec.add_run(
                "CRITICAL WARNING: The region demonstrates high vulnerability to hydrological emergencies. "
                "It is strongly recommended to restrict permanent concrete civil structures in buffers closer than 500m "
                "to surface channels. Implement permeable pavement surfaces, sustainable drainage channels, "
                "and deploy real-time telemetry river monitors to build resilient local infrastructure."
            )
            p_rec_run.bold = True
            p_rec_run.font.color.rgb = docx.shared.RGBColor(153, 27, 27) # Dark Red
        else:
            p_rec.add_run(
                "The analysis indicates low-to-moderate environmental hazards. Standard structural guidelines are "
                "adequate. Urban growth should prioritize preserving green pathways and planting vegetation strips "
                "to prevent heat island effects and retain historical natural absorption thresholds."
            )
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

# Global Instance
report_generator = ReportGeneratorService()
